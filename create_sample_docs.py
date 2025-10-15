"""
Create Sample Banking Documents for RAG System

WHY THIS SCRIPT:
- Creates realistic banking policy documents
- These documents will be indexed by our RAG system
- Agent will retrieve information from these to answer questions

LEARNING POINTS:
- Creating structured documents programmatically
- Banking domain knowledge
- Document organization for RAG
"""

from pathlib import Path
from docx import Document
from datetime import datetime

# Create output directory
output_dir = Path("data/sample_banking_docs")
output_dir.mkdir(parents=True, exist_ok=True)

print("\n🏦 Creating Sample Banking Documents for RAG System...\n")


def create_account_types_doc():
    """Create document about account types and features"""
    doc = Document()
    doc.add_heading('DemoBank Account Types and Features', 0)
    doc.add_paragraph(f'Last Updated: {datetime.now().strftime("%B %Y")} | Version 2.1')
    
    # Savings Accounts
    doc.add_heading('Savings Accounts', 1)
    
    doc.add_heading('Standard Savings Account', 2)
    doc.add_paragraph(
        'Our Standard Savings Account is perfect for building your emergency fund or saving for short-term goals. '
        'Key features include:'
    )
    doc.add_paragraph('• Interest Rate: 2.5% APY on balances over $1,000', style='List Bullet')
    doc.add_paragraph('• Interest Rate: 2.0% APY on balances under $1,000', style='List Bullet')
    doc.add_paragraph('• Minimum Opening Deposit: $100', style='List Bullet')
    doc.add_paragraph('• No Monthly Maintenance Fee', style='List Bullet')
    doc.add_paragraph('• Unlimited free ATM withdrawals at DemoBank ATMs', style='List Bullet')
    doc.add_paragraph('• FDIC insured up to $250,000', style='List Bullet')
    doc.add_paragraph('• Online and mobile banking access included', style='List Bullet')
    
    doc.add_heading('Premium Savings Account', 2)
    doc.add_paragraph(
        'For customers with higher balances, our Premium Savings Account provides enhanced benefits and higher rates:'
    )
    doc.add_paragraph('• Interest Rate: 3.2% APY on balances over $10,000', style='List Bullet')
    doc.add_paragraph('• Interest Rate: 2.8% APY on balances $5,000-$9,999', style='List Bullet')
    doc.add_paragraph('• Interest Rate: 2.5% APY on balances under $5,000', style='List Bullet')
    doc.add_paragraph('• Minimum Opening Deposit: $5,000', style='List Bullet')
    doc.add_paragraph('• No Monthly Fee with minimum balance of $5,000', style='List Bullet')
    doc.add_paragraph('• $10 monthly fee if balance falls below $5,000', style='List Bullet')
    doc.add_paragraph('• Priority customer service access', style='List Bullet')
    doc.add_paragraph('• Free cashier\'s checks and money orders', style='List Bullet')
    
    # Checking Accounts
    doc.add_heading('Checking Accounts', 1)
    
    doc.add_heading('Essential Checking Account', 2)
    doc.add_paragraph('Designed for everyday banking needs with no minimum balance requirement:')
    doc.add_paragraph('• No Minimum Balance Required', style='List Bullet')
    doc.add_paragraph('• Monthly Fee: $8', style='List Bullet')
    doc.add_paragraph('• Fee Waived with: $500 minimum daily balance OR one direct deposit per month', style='List Bullet')
    doc.add_paragraph('• Free debit card with contactless payment', style='List Bullet')
    doc.add_paragraph('• Unlimited check writing', style='List Bullet')
    doc.add_paragraph('• Overdraft protection available for $10/transfer', style='List Bullet')
    doc.add_paragraph('• Mobile check deposit enabled', style='List Bullet')
    
    doc.save(output_dir / "account_types.docx")
    print("✅ Created: account_types.docx")


def create_fees_schedule_doc():
    """Create fee schedule document"""
    doc = Document()
    doc.add_heading('DemoBank Fee Schedule', 0)
    doc.add_paragraph(f'Effective Date: {datetime.now().strftime("%B %Y")}')
    
    doc.add_heading('Monthly Account Fees', 1)
    doc.add_paragraph('Essential Checking: $8 per month (waived with $500 min balance or direct deposit)', style='List Bullet')
    doc.add_paragraph('Premium Savings: $10 per month (waived with $5,000 min balance)', style='List Bullet')
    doc.add_paragraph('Standard Savings: No monthly fee', style='List Bullet')
    doc.add_paragraph('Business Checking: $15 per month (waived with $2,500 min balance)', style='List Bullet')
    
    doc.add_heading('Transaction Fees', 1)
    doc.add_paragraph('Overdraft Fee: $35 per transaction (maximum 3 per day)', style='List Bullet')
    doc.add_paragraph('NSF (Non-Sufficient Funds) Fee: $35 per item', style='List Bullet')
    doc.add_paragraph('Stop Payment Fee: $30 per request', style='List Bullet')
    doc.add_paragraph('Paper Statement Fee: $3 per statement (e-statements are free)', style='List Bullet')
    doc.add_paragraph('Wire Transfer Domestic: $25 outgoing, $15 incoming', style='List Bullet')
    doc.add_paragraph('Wire Transfer International: $45 outgoing, $20 incoming', style='List Bullet')
    
    doc.add_heading('ATM Fees', 1)
    doc.add_paragraph('DemoBank ATM: Free unlimited transactions', style='List Bullet')
    doc.add_paragraph('Non-DemoBank ATM (U.S.): $3 per transaction', style='List Bullet')
    doc.add_paragraph('International ATM: $5 per transaction plus 3% foreign transaction fee', style='List Bullet')
    doc.add_paragraph('Premium account holders: Up to $20 in ATM fee rebates monthly', style='List Bullet')
    
    doc.save(output_dir / "fee_schedule.docx")
    print("✅ Created: fee_schedule.docx")


def create_loan_products_doc():
    """Create loan products document"""
    doc = Document()
    doc.add_heading('DemoBank Loan Products', 0)
    
    doc.add_heading('Personal Loans', 1)
    doc.add_paragraph('Flexible financing for your personal needs:')
    doc.add_paragraph('• Loan Amount: $1,000 to $50,000', style='List Bullet')
    doc.add_paragraph('• Interest Rate: 6.99% - 18.99% APR (based on creditworthiness)', style='List Bullet')
    doc.add_paragraph('• Repayment Terms: 12 to 60 months', style='List Bullet')
    doc.add_paragraph('• No prepayment penalties', style='List Bullet')
    doc.add_paragraph('• Fixed monthly payments', style='List Bullet')
    doc.add_paragraph('• Funding within 1-2 business days upon approval', style='List Bullet')
    
    doc.add_paragraph('\nEligibility Requirements:')
    doc.add_paragraph('• Minimum credit score: 650', style='List Bullet')
    doc.add_paragraph('• Annual income: $30,000 minimum', style='List Bullet')
    doc.add_paragraph('• Debt-to-income ratio: Maximum 43%', style='List Bullet')
    doc.add_paragraph('• At least 18 years old with U.S. citizenship or permanent residency', style='List Bullet')
    
    doc.add_heading('Auto Loans', 1)
    doc.add_paragraph('Finance your vehicle with competitive rates:')
    doc.add_paragraph('• New Cars: 5.49% - 8.99% APR', style='List Bullet')
    doc.add_paragraph('• Used Cars (up to 5 years old): 6.49% - 10.99% APR', style='List Bullet')
    doc.add_paragraph('• Loan Terms: 24 to 72 months', style='List Bullet')
    doc.add_paragraph('• Maximum Loan Amount: $75,000', style='List Bullet')
    doc.add_paragraph('• No application fee', style='List Bullet')
    doc.add_paragraph('• Pre-approval available online', style='List Bullet')
    
    doc.add_heading('Home Mortgages', 1)
    doc.add_paragraph('Purchase your dream home:')
    doc.add_paragraph('• 30-Year Fixed Rate: 6.75% APR', style='List Bullet')
    doc.add_paragraph('• 15-Year Fixed Rate: 6.00% APR', style='List Bullet')
    doc.add_paragraph('• 5/1 ARM: Starting at 5.875% APR', style='List Bullet')
    doc.add_paragraph('• Down Payment: As low as 3% for qualified borrowers', style='List Bullet')
    doc.add_paragraph('• Maximum Loan Amount: $2,000,000', style='List Bullet')
    doc.add_paragraph('• First-time homebuyer programs available', style='List Bullet')
    
    doc.save(output_dir / "loan_products.docx")
    print("✅ Created: loan_products.docx")


def create_customer_service_doc():
    """Create customer service information document"""
    doc = Document()
    doc.add_heading('DemoBank Customer Service Guide', 0)
    
    doc.add_heading('Contact Information', 1)
    doc.add_paragraph('24/7 Customer Support:')
    doc.add_paragraph('• Phone: 1-800-DEMO-BANK (1-800-336-6226)', style='List Bullet')
    doc.add_paragraph('• International: +1-555-123-4567', style='List Bullet')
    doc.add_paragraph('• Email: support@demobank.com', style='List Bullet')
    doc.add_paragraph('• Live Chat: Available through mobile app and website 24/7', style='List Bullet')
    doc.add_paragraph('• Mail: DemoBank Customer Service, PO Box 12345, New York, NY 10001', style='List Bullet')
    
    doc.add_heading('Branch Hours', 1)
    doc.add_paragraph('Standard Branch Hours:')
    doc.add_paragraph('• Monday - Friday: 9:00 AM - 5:00 PM', style='List Bullet')
    doc.add_paragraph('• Saturday: 9:00 AM - 2:00 PM', style='List Bullet')
    doc.add_paragraph('• Sunday: Closed', style='List Bullet')
    doc.add_paragraph('• Extended hours available at select locations', style='List Bullet')
    
    doc.add_heading('Common Services', 1)
    doc.add_paragraph('Account Opening: Visit any branch or apply online in 10 minutes', style='List Bullet')
    doc.add_paragraph('Card Replacement: Order through mobile app, arrives in 5-7 business days', style='List Bullet')
    doc.add_paragraph('Dispute Resolution: Submit disputes within 60 days of transaction', style='List Bullet')
    doc.add_paragraph('Password Reset: Use "Forgot Password" link or call customer support', style='List Bullet')
    
    doc.save(output_dir / "customer_service.docx")
    print("✅ Created: customer_service.docx")


def create_security_policy_doc():
    """Create security and fraud prevention document"""
    doc = Document()
    doc.add_heading('DemoBank Security & Fraud Prevention', 0)
    
    doc.add_heading('Fraud Detection', 1)
    doc.add_paragraph(
        'DemoBank employs advanced AI-powered fraud detection that monitors your accounts 24/7. '
        'We analyze transaction patterns, geographic locations, and spending behavior to identify '
        'suspicious activity before it affects your account.'
    )
    
    doc.add_heading('Zero Liability Protection', 1)
    doc.add_paragraph('You are protected from unauthorized transactions when you report them promptly:')
    doc.add_paragraph('• Report within 2 business days: $0 liability', style='List Bullet')
    doc.add_paragraph('• Report within 60 days: Maximum $50 liability', style='List Bullet')
    doc.add_paragraph('• Report fraud immediately: 1-800-DEMO-BANK', style='List Bullet')
    
    doc.add_heading('Card Controls', 1)
    doc.add_paragraph('Through our mobile app, you can:')
    doc.add_paragraph('• Instantly lock/unlock your debit card', style='List Bullet')
    doc.add_paragraph('• Set spending limits by merchant category', style='List Bullet')
    doc.add_paragraph('• Block international transactions', style='List Bullet')
    doc.add_paragraph('• Disable ATM withdrawals temporarily', style='List Bullet')
    doc.add_paragraph('• Receive instant transaction alerts', style='List Bullet')
    
    doc.save(output_dir / "security_policy.docx")
    print("✅ Created: security_policy.docx")


# Create all documents
if __name__ == "__main__":
    create_account_types_doc()
    create_fees_schedule_doc()
    create_loan_products_doc()
    create_customer_service_doc()
    create_security_policy_doc()
    
    print(f"\n✅ All documents created in: {output_dir}")
    print(f"📊 Total documents: {len(list(output_dir.glob('*.docx')))}")
    print("\n💡 These documents will be used by the RAG system to answer customer questions accurately.\n")